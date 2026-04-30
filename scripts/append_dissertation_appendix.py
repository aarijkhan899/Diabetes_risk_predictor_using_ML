#!/usr/bin/env python3
"""
Append a professionally formatted appendix (code listings + verified sample outputs)
to the combined dissertation DOCX.

Usage:
  python scripts/append_dissertation_appendix.py "Aariz_Final_Combined_Dissertation (1).docx"
"""

from __future__ import annotations

import argparse
import subprocess
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
except ImportError as e:
    print("Requires python-docx: pip install python-docx", file=sys.stderr)
    raise SystemExit(1) from e

ROOT = Path(__file__).resolve().parents[1]

# Keep Word comfortable: fewer than ~50 lines per monospaced paragraph avoids huge paragraphs.
_CHUNK_LINES = 42


def _paragraph_shading_light_gray(paragraph) -> None:
    """Soft fill behind listings (readable on screen and print PDF)."""
    p_pr = paragraph._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F5F5")
    shd.set(qn("w:val"), "clear")
    p_pr.append(shd)


def listing_caption(doc: Document, repo_relative_path: str, line_start_1based: int, line_end_1based: int) -> None:
    para = doc.add_paragraph()
    run = para.add_run(f"Listing — {repo_relative_path} (lines {line_start_1based}-{line_end_1based})")
    run.italic = True
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(2)


def monospace_block(doc: Document, body: str, *, shaded: bool = True) -> None:
    raw = body.strip("\n").split("\n")
    chunks: list[list[str]] = []
    buf: list[str] = []
    for ln in raw:
        buf.append(ln)
        if len(buf) >= _CHUNK_LINES:
            chunks.append(buf)
            buf = []
    if buf:
        chunks.append(buf)
    first = True
    for ci, blk in enumerate(chunks):
        last_blk = ci == len(chunks) - 1
        para = doc.add_paragraph()
        fmt = para.paragraph_format
        fmt.left_indent = Inches(0.25)
        fmt.right_indent = Inches(0.12)
        fmt.space_before = Pt(2 if not first else 4)
        fmt.space_after = Pt(10 if last_blk else 2)
        if shaded:
            _paragraph_shading_light_gray(para)
        run = para.add_run("\n".join(blk))
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        first = False


def prose(doc: Document, text: str) -> None:
    para = doc.add_paragraph(text.replace("\u2019", "'"))
    para.paragraph_format.space_after = Pt(8)


def bold_title(doc: Document, text: str, *, size_pt: int = 14) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)


def collect_sample_json(repo_root: Path) -> tuple[str, str, str]:
    """Run Flask test_client; return meta file, GET /health, POST /predict JSON."""
    snippet = '''
import logging, json, importlib.util
from pathlib import Path

logging.disable(logging.CRITICAL)
spec = importlib.util.spec_from_file_location("api", Path("ml") / "api.py")
mod = importlib.util.module_from_spec(spec)
spec.loader.exec_module(mod)
app = mod.app
c = app.test_client()

meta_txt = Path("ml/models/model_meta.json").read_text(encoding="utf-8")
health = json.dumps(c.get("/health").get_json(), indent=2)

body = {
    "Pregnancies": 6,
    "Glucose": 148,
    "BloodPressure": 72,
    "SkinThickness": 35,
    "Insulin": 0,
    "BMI": 33.6,
    "DiabetesPedigreeFunction": 0.627,
    "Age": 50,
}
pred_j = json.dumps(c.post("/predict", json=body).get_json(), indent=2)
print("__META_START__")
print(meta_txt.strip())
print("__META_END__")
print("__HEALTH_START__")
print(health.strip())
print("__HEALTH_END__")
print("__PREDICT_START__")
print(pred_j.strip())
print("__PREDICT_END__")
'''
    proc = subprocess.run(
        [sys.executable, "-c", snippet],
        cwd=str(repo_root),
        capture_output=True,
        text=True,
        encoding="utf-8",
        timeout=120,
    )
    out = proc.stdout
    err = proc.stderr
    if proc.returncode != 0:
        raise RuntimeError(f"Sample run failed ({proc.returncode}): {err or out}")

    def block(tag: str) -> str:
        a = out.index(f"__{tag}_START__") + len(f"__{tag}_START__")
        b = out.index(f"__{tag}_END__")
        return out[a:b].strip()

    return block("META"), block("HEALTH"), block("PREDICT")


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("docx", nargs="?", default=str(ROOT / "Aariz_Final_Combined_Dissertation (1).docx"))
    args = ap.parse_args()
    path = Path(args.docx).resolve()
    if not path.is_file():
        raise SystemExit(f"File not found: {path}")

    doc = Document(str(path))

    if any(
        "appendix — implementation listing" in (p.text or "").lower()
        or "implementation listing & sample artefacts" in (p.text or "").lower()
        for p in doc.paragraphs
    ):
        raise SystemExit(
            "Appendix already present in this document (refusing to duplicate). "
            "Restore from the *_backup_before_appendix.docx copy if needed, then edit and re-run."
        )

    meta_js, health_js, predict_js = collect_sample_json(ROOT)

    tp = ROOT / "ml/train_model.py"
    tl = tp.read_text(encoding="utf-8").splitlines()
    train_slice_from = 100  # 0-based exclusive start before `def preprocess` (line 101)
    train_slice_to = min(173, len(tl))
    train_slice = tl[train_slice_from:train_slice_to]
    train_code = "\n".join(train_slice).rstrip() + "\n\n# ... Evaluation, artefact persistence, main() trimmed ..."
    train_caption_hi = train_slice_from + len(train_slice)  # exclusive end idx -> line numbers up to train_caption_hi

    api_p = ROOT / "ml/api.py"
    al = api_p.read_text(encoding="utf-8").splitlines()
    api_end = min(200, len(al))
    api_code = "\n".join(al[20:api_end]).strip()
    api_code += "\n\n# ... Flask application entry trimmed ..."

    rails_path = ROOT / "diabetes_app/app/controllers/predictions_controller.rb"
    rl = rails_path.read_text(encoding="utf-8").splitlines()
    rails_end = min(94, len(rl))
    rails_code = (
        "\n".join(rl[:rails_end]).rstrip()
        + "\n\n# ... #about, validation helpers, remainder omitted ..."
    )

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    bold_title(doc, "Appendix — Implementation listing & sample artefacts", size_pt=16)

    prose(
        doc,
        "This appendix documents the artefact-aligned implementation underpinning Chapters discussing system design "
        "(Python preprocessing and classifier selection; Flask inference service; Rails web tier). Source listings are verbatim from the Git "
        "repository, captions give repository-relative paths plus line ranges (numbered from 1 within each file). Monospaced text uses Courier New; "
        "JSON fragments were emitted by Flask's bundled test_client after loading persisted joblib models (figures update if you re-run training).",
    )

    bold_title(doc, "Appendix A — Preprocessing & model selection excerpt", size_pt=12)

    prose(
        doc,
        "Imputation for biologically implausible sentinel zeros on selected physiological fields; StandardScaler; SMOTE; stratified ROC-AUC "
        "GridSearchCV across heterogeneous estimators (see repository for hyperparameter grids and evaluation persistence).",
    )
    listing_caption(doc, "ml/train_model.py", train_slice_from + 1, train_caption_hi)
    monospace_block(doc, train_code)

    bold_title(doc, "Appendix B — REST API: artefacts, HTTP surface, inference", size_pt=12)

    prose(
        doc,
        "Feature vocabulary, persistence paths, explanatory SHAP aggregation for tree ensembles with multi-output SHAP tensors, probabilistic inference, JSON contract incl. disclaimers.",
    )
    listing_caption(doc, "ml/api.py", 21, api_end)
    monospace_block(doc, api_code)

    bold_title(doc, "Appendix C — Rails orchestration + HTTParty client", size_pt=12)

    prose(
        doc,
        "Validated user input mirrored against clinical min/max envelopes; synchronous JSON POST toward ML_API_URL; timeouts and surfaced error taxonomy.",
    )
    listing_caption(doc, "diabetes_app/app/controllers/predictions_controller.rb", 1, rails_end)
    monospace_block(doc, rails_code)

    bold_title(doc, "Appendix D — Persisted classifier metadata snapshot", size_pt=12)

    prose(doc, "Captured ml/models/model_meta.json describing the artefact pinned at appendix generation.")
    listing_caption(doc, "ml/models/model_meta.json", 1, meta_js.strip().count("\n") + 1)
    monospace_block(doc, meta_js)

    bold_title(doc, "Appendix E — Sample inference payloads", size_pt=12)

    prose(
        doc,
        "Automated reproducibility hook: Flask test_client, GET /health and POST /predict with the canonical Rails EXAMPLE_PATIENT vector.",
    )
    listing_caption(doc, "Observed REST responses", 1, predict_js.strip().count("\n") + 8)
    monospace_block(
        doc,
        "[GET /health]\n" + "=" * 72 + "\n" + health_js + "\n\n[POST /predict]\n" + "=" * 72 + "\n" + predict_js,
        shaded=False,
    )

    doc.save(str(path))
    print(f"Saved: {path}")


if __name__ == "__main__":
    main()
